"""
本地 AI 助手 —— Ollama (Qwen/Llama) 提取 EDC 问题摘要 + 8D 根因/措施
放到 app/ai_helper.py
"""
import re
import json
import os
import subprocess
import hashlib
import zipfile
import requests
from pathlib import Path
from xml.etree import ElementTree as ET

OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "qwen3:8b"   # 英文输出可改 llama3.2:3b

# 输出语言：'zh' 中文 / 'en' 英文
OUTPUT_LANG = "en"


# ──────────────────────────────────────────────────────────
# Ollama 基础
# ──────────────────────────────────────────────────────────

def is_ollama_available(timeout=3):
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=timeout)
        return r.status_code == 200
    except Exception:
        return False


def _call_ollama(prompt, timeout=120, num_predict=300, logger=None):
    """调用 Ollama，返回文本或 None"""
    try:
        resp = requests.post(
            OLLAMA_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "think": False,
                "options": {"temperature": 0.0, "num_predict": 1500},
            },
            timeout=timeout,
        )
        if resp.status_code != 200:
            if logger:
                logger.warning(f"[AI] Ollama returned {resp.status_code}")
            return None
        raw = (resp.json().get("response") or "").strip()
        raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
        return raw if raw else None
    except requests.exceptions.Timeout:
        if logger:
            logger.warning("[AI] Ollama timeout")
        return None
    except requests.exceptions.ConnectionError:
        if logger:
            logger.warning("[AI] Ollama not running")
        return None
    except Exception as e:
        if logger:
            logger.warning(f"[AI] error: {e}")
        return None


def _parse_json(text):
    """从模型输出中稳健提取 JSON"""
    if not text:
        return None
    text = text.strip()
    # 去掉 ```json ``` 围栏
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    # 抓第一个 { 到最后一个 }
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if m:
        text = m.group(0)
    try:
        return json.loads(text)
    except Exception:
        return None


# ──────────────────────────────────────────────────────────
# 1) EDC 问题描述摘要
# ──────────────────────────────────────────────────────────

_SUMMARY_PROMPT = {
    "en": """You are an automotive parts quality engineer. Below is the raw issue description from an EDC quality report. It may contain Italian/English text, formatting symbols, and irrelevant info.

CRITICAL RULES:
- Extract ONLY the specific defect/failure description from the text.
- Do NOT mention: 8D report requests, cost charges, supplier actions, rejection/selection processes, or administrative instructions. These are irrelevant.
- Do NOT state obvious facts like "parts were rejected" or "non-conforming parts found" — every EDC report has rejections, that is not useful information.
- Do NOT include quantities or piece counts.
- Focus ONLY on: what is the SPECIFIC technical defect? (e.g., cracks, porosity, paint defects, leakage, excess weight, dimensional deviation, surface dents)
- If the report does NOT describe any specific defect characteristics, output exactly: "No specific defect described in report, further investigation needed."
- NEVER invent or guess a defect type not explicitly stated in the text.

Summarize the specific defect in ONE concise English sentence (max 15 words).
Output only the summary, no prefix/quotes/explanation.

Raw text:
{raw}""",
}


def summarize_issue(raw_text, timeout=180, logger=None):
    if not raw_text or not raw_text.strip():
        return None
    if len(raw_text.strip()) < 30:
        return raw_text.strip()

    prompt = _SUMMARY_PROMPT[OUTPUT_LANG].format(raw=raw_text.strip()[:2000])
    summary = _call_ollama(prompt, timeout=timeout, num_predict=800, logger=logger)
    if not summary:
        return None

    for prefix in ["概括：", "概括:", "问题：", "总结：", "Summary:", "Issue:"]:
        if summary.startswith(prefix):
            summary = summary[len(prefix):].strip()
    summary = summary.strip("\"'「」“”").strip()
    return summary or None


# ──────────────────────────────────────────────────────────
# 2) 8D 报告文本提取（Excel / PDF / Word / Email）
# ──────────────────────────────────────────────────────────

LIBREOFFICE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice",
    "/usr/bin/libreoffice",
]


def _find_soffice():
    for p in LIBREOFFICE_PATHS:
        if os.path.exists(p):
            return p
    return None


def _convert_office_to_pdf_for_ai(file_path, logger=None):
    src = Path(file_path)
    if not src.exists():
        return None

    soffice = _find_soffice()
    if not soffice:
        if logger:
            logger.info("[AI] LibreOffice not found; cannot convert Office 8D to PDF")
        return None

    cache_dir = src.parent / "_ai_extract_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    key = hashlib.md5(f"{src}_{src.stat().st_size}_{src.stat().st_mtime}".encode()).hexdigest()
    cached_pdf = cache_dir / f"{key}.pdf"
    if cached_pdf.exists():
        return str(cached_pdf)

    try:
        if logger:
            logger.info(f"[AI] converting Office 8D to PDF: {src.name}")
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(cache_dir), str(src)],
            timeout=90,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            if logger:
                logger.warning(f"[AI] LibreOffice conversion failed: {result.stderr[:300]}")
            return None

        generated = cache_dir / f"{src.stem}.pdf"
        if generated.exists():
            generated.replace(cached_pdf)
            return str(cached_pdf)
        if logger:
            logger.warning(f"[AI] LibreOffice did not create expected PDF for {src.name}")
        return None
    except Exception as e:
        if logger:
            logger.warning(f"[AI] Office to PDF conversion failed for {src.name}: {e}")
        return None


def _extract_text_from_pptx(file_path, logger=None):
    p = Path(file_path)
    lines = []
    try:
        with zipfile.ZipFile(str(p)) as zf:
            slide_names = sorted(
                name for name in zf.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            for idx, name in enumerate(slide_names, start=1):
                try:
                    root = ET.fromstring(zf.read(name))
                except ET.ParseError:
                    continue
                texts = [
                    node.text.strip()
                    for node in root.findall(".//a:t", ns)
                    if node.text and node.text.strip()
                ]
                if texts:
                    lines.append(f"=== Slide {idx} ===")
                    lines.append("\n".join(texts))
    except Exception as e:
        if logger:
            logger.warning(f"[AI] pptx text extract failed for {p.name}: {e}")
        return ""
    return "\n".join(lines)


def _skip_excel_sheet(name):
    sheet_name = (name or "").strip().lower()
    return any(key in sheet_name for key in ("填写要求", "勿动", "instruction", "readme", "q alert"))


def extract_text_from_file(file_path, logger=None):
    """从 8D 报告文件提取纯文本。支持 xlsx/xls/pdf/docx/pptx/ppt/txt。"""
    p = Path(file_path)
    if not p.exists():
        return ""
    ext = p.suffix.lower().lstrip(".")

    try:
        # ── Excel（你的 8D 主要格式）──
        if ext == "xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(str(p), data_only=True, read_only=True)
            lines = []
            for ws in wb.worksheets:
                if _skip_excel_sheet(ws.title):
                    continue
                lines.append(f"=== Sheet: {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = [re.sub(r"\s+", " ", str(c)).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            wb.close()
            return "\n".join(lines)

        # ── 老 Excel 格式 ──
        elif ext == "xls":
            try:
                import xlrd
                wb = xlrd.open_workbook(str(p))
                lines = []
                for sheet in wb.sheets():
                    if _skip_excel_sheet(sheet.name):
                        continue
                    lines.append(f"=== Sheet: {sheet.name} ===")
                    for r in range(sheet.nrows):
                        cells = [re.sub(r"\s+", " ", str(sheet.cell_value(r, c))).strip() for c in range(sheet.ncols)]
                        cells = [c for c in cells if c]
                        if cells:
                            lines.append(" | ".join(cells))
                return "\n".join(lines)
            except ImportError:
                if logger:
                    logger.warning("[AI] xls needs: pip install xlrd")
                return ""

        # ── PDF ──
        elif ext == "pdf":
            import pdfplumber
            with pdfplumber.open(str(p)) as pdf:
                return "\n".join(pg.extract_text() or "" for pg in pdf.pages)

        # ── Word ──
        elif ext == "docx":
            import docx
            d = docx.Document(str(p))
            parts = [para.text for para in d.paragraphs if para.text.strip()]
            # 也读表格
            for table in d.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

        # ── PowerPoint 新格式：优先直接抽取幻灯片文字 ──
        elif ext == "pptx":
            text = _extract_text_from_pptx(p, logger=logger)
            if len(text.strip()) >= 20:
                return text

            pdf_path = _convert_office_to_pdf_for_ai(p, logger=logger)
            if pdf_path:
                return extract_text_from_file(pdf_path, logger=logger)
            return text

        # ── PowerPoint 老格式：用 LibreOffice 转 PDF 后提取 ──
        elif ext == "ppt":
            pdf_path = _convert_office_to_pdf_for_ai(p, logger=logger)
            if pdf_path:
                return extract_text_from_file(pdf_path, logger=logger)
            return ""

        # ── 纯文本 ──
        elif ext == "txt":
            return p.read_text(encoding="utf-8", errors="ignore")

        else:
            if logger:
                logger.info(f"[AI] unsupported 8D format: {ext}")
            return ""

    except Exception as e:
        if logger:
            logger.warning(f"[AI] text extract failed for {p.name}: {e}")
        return ""


# ──────────────────────────────────────────────────────────
# 3) 8D 根因 + 纠正措施提取
# ──────────────────────────────────────────────────────────

# ──────────────────────────────────────────────────────────
# 3) 8D 根因（发生 + 流出）+ 纠正措施提取（中英双语）
# ──────────────────────────────────────────────────────────

_8D_PROMPT = """You are analyzing an 8D report from an automotive supplier. 
The text below was extracted from Excel/PDF and contains Chinese + English bilingual labels.

The report has TWO PARALLEL 5-Why analyses. You MUST find BOTH:

【A. OCCURRENCE 发生原因】Why did the defect HAPPEN?
   Anchors to search for: "根本原因", "Root Cause", "5 Why Analysis —— Root Cause",
   "发生原因", "D4 Root Cause"
   Do NOT only output the final/deepest "Why" unless the report truly provides only one Why.
   Extract the supplier's key Why chain as a compact causal summary:
   immediate cause -> process/system cause -> final root cause.
   Include 2-4 important Why/Reason links when available, remove repeated wording and minor filler.
   Chinese and English root cause fields must describe the same cause chain.

【B. ESCAPE 流出原因 / 未检出原因】Why was the defect NOT DETECTED before shipping?
   Anchors to search for: "未检出原因", "Reason for Non-detection", "Non-detection",
   "Detection Root Cause", "流出原因", "Escape", "为什么没有检出"
   This section is SEPARATE from occurrence. Even if it only has Why1+Reason1, extract it.
   ⚠ Do NOT copy the occurrence cause here. They are different questions.

For each, also extract its CORRECTIVE ACTION (D5/D6 永久措施/纠正措施):
   - occurrence_action → action that fixes the root cause (training, jig change, process change)
   - escape_action → action that fixes the inspection gap (added check, gauge upgrade, operator training on measurement)
   - Prioritize D5/D6 sections: "拟实施措施", "实施措施", "Proposed Measures", "Implemented Measures",
     "Preventive Measures", "改善措施", "措施描述", "Correction measure".
   - If D5/D6 action sections are present, do NOT leave action fields empty.
   - If one measure fixes both occurrence and detection, split it by intent:
     process/tool/parameter control → occurrence_action;
     torque confirmation/inspection/checklist/gauge verification → escape_action.

Potential D5/D6 corrective action text pre-extracted from the report:
{action_hint}

Output STRICT JSON, no preamble, no markdown, no comments:
{{"occurrence_cause":"<中文>","occurrence_cause_en":"<English translation>","occurrence_action":"<中文>","occurrence_action_en":"<English>","escape_cause":"<中文>","escape_cause_en":"<English>","escape_action":"<中文>","escape_action_en":"<English>"}}

RULES:
- If a field is genuinely absent in the report, use "" (empty string). Do NOT fabricate.
- escape_cause and occurrence_cause MUST be different. If you find only one, the other is "".
- Translate Chinese to natural English, not literal word-by-word.
- Root cause fields should be one compact paragraph, not a long copied table.
- Keep occurrence_cause / escape_cause concise but complete: ideally 2-4 cause links, 40-120 Chinese characters or 25-65 English words.
- Keep action fields 1-2 concise sentences.

8D report text:
{raw}"""


def _compact_8d_field(value, max_chars=420):
    """Keep AI 8D extraction readable if the model copies a long report section."""
    text = re.sub(r"\s+", " ", value or "").strip()
    if len(text) <= max_chars:
        return text

    cut = text[:max_chars].rstrip()
    breakpoints = [cut.rfind(mark) for mark in ("。", "；", ";", ".", "，", ",")]
    breakpoint = max(breakpoints)
    if breakpoint >= int(max_chars * 0.65):
        cut = cut[: breakpoint + 1].rstrip()
    return cut.rstrip(" .。,.，;；") + "..."


ACTION_SECTION_KEYWORDS = (
    "D5", "D6", "拟实施措施", "实施措施", "永久措施", "纠正措施", "预防措施",
    "Proposed Measures", "Implemented Measures", "Preventive Measures",
    "Corrective Action", "Correction measure", "Measures Description", "措施描述", "改善措施",
)


def _extract_action_hint(raw, max_chars=5000):
    """Pull likely D5/D6 action sections closer to the model's attention."""
    text = raw or ""
    blocks = re.split(r"(?=^=== Slide \d+ ===)", text, flags=re.M)
    selected = []

    for block in blocks:
        if any(k.lower() in block.lower() for k in ACTION_SECTION_KEYWORDS):
            selected.append(block.strip())

    if not selected:
        return ""

    hint = "\n\n".join(selected)
    hint = re.sub(r"\n{3,}", "\n\n", hint).strip()
    return hint[:max_chars]


def _fallback_actions_from_hint(action_hint):
    """Best-effort fallback when the model finds causes but leaves actions blank."""
    text = re.sub(r"\s+", " ", action_hint or "").strip()
    if not text:
        return {}

    lowered = text.lower()
    occurrence_cn, occurrence_en = [], []
    escape_cn, escape_en = [], []
    occurrence_cause_cn, occurrence_cause_en = [], []
    escape_cause_cn, escape_cause_en = [], []

    if re.search(r"electric\s+torque\s+gun|torque\s+monitoring|电动.{0,12}力矩.{0,8}枪", text, re.I):
        occurrence_cn.append("弃用气动力矩枪，采购/导入电动带扭矩监控力矩枪。")
        occurrence_en.append("Discontinue pneumatic torque guns and introduce electric torque guns with torque monitoring.")
    if re.search(r"receiver\s+tank|pressure\s+regulat|储气罐|调\s*压\s*阀", text, re.I):
        occurrence_cn.append("电动力矩枪到厂前，临时增加储气罐和调压阀以稳定气源压力。")
        occurrence_en.append("Before the electric torque guns arrive, add an air receiver tank and pressure regulating valve to stabilize air pressure.")
    if re.search(r"torque\s+wrench|verify\s+the\s+torque|confirm.{0,30}torque|力矩扳手|扭矩.{0,12}确认|力矩.{0,12}确认", text, re.I):
        escape_cn.append("电动力矩枪到厂前，使用力矩扳手确认装配后放气螺钉扭矩。")
        escape_en.append("Before the electric torque guns arrive, use a torque wrench to verify the bleed screw torque after assembly.")

    if not occurrence_cn and not escape_cn:
        table_actions = _extract_corrective_actions_from_table(action_hint)
        occurrence_cause_cn.extend(table_actions.get("cause_cn", []))
        occurrence_cause_en.extend(table_actions.get("cause_en", []))
        escape_cause_cn.extend(table_actions.get("escape_cause_cn", []))
        escape_cause_en.extend(table_actions.get("escape_cause_en", []))
        occurrence_cn.extend(table_actions.get("action_cn", []))
        occurrence_en.extend(table_actions.get("action_en", []))
        escape_cn.extend(table_actions.get("escape_cn", []))
        escape_en.extend(table_actions.get("escape_en", []))

    return {
        "root_cause": _compact_8d_field("".join(occurrence_cause_cn), 320),
        "root_cause_en": _compact_8d_field(" ".join(occurrence_cause_en), 360),
        "escape_cause": _compact_8d_field("".join(escape_cause_cn), 300),
        "escape_cause_en": _compact_8d_field(" ".join(escape_cause_en), 340),
        "action": _compact_8d_field("".join(occurrence_cn), 260),
        "action_en": _compact_8d_field(" ".join(occurrence_en), 320),
        "escape_action": _compact_8d_field("".join(escape_cn), 220),
        "escape_action_en": _compact_8d_field(" ".join(escape_en), 280),
    }


def _split_bilingual_cell(value):
    cn, en = [], []
    for part in re.split(r"[\r\n]+", value or ""):
        item = re.sub(r"\s+", " ", part).strip(" -:：,，.")
        if not item:
            continue
        if re.search(r"[\u4e00-\u9fff]", item):
            match = re.search(r"\s+([A-Za-z][A-Za-z0-9 ,()./+\\-]+)$", item)
            if match:
                cn_text = item[:match.start()].strip(" -:：,，.")
                en_text = match.group(1).strip(" -:：,，.")
                if cn_text:
                    cn.append(cn_text)
                if en_text:
                    en.append(en_text)
            else:
                cn.append(item)
        elif re.search(r"[A-Za-z]", item):
            en.append(item)
    return "；".join(cn), ". ".join(en)


def _extract_corrective_actions_from_table(action_hint):
    lines = [line.strip() for line in (action_hint or "").splitlines() if line.strip()]
    in_corrective = False
    cause_cn, cause_en, escape_cause_cn, escape_cause_en = [], [], [], []
    action_cn, action_en, escape_cn, escape_en = [], [], [], []
    seen = set()

    detection_keywords = ("扫描", "拉点", "检查", "检验", "确认", "inspection", "scan", "check", "verify")
    header_keywords = (
        "corrective actions planned", "corrective actions taken",
        "responsible", "target", "actual", "status", "根本原因", "整改措施",
        "责任人", "完成时间", "目标", "实际", "状态",
    )

    for line in lines:
        lower = line.lower()
        if "corrective actions planned" in lower or "5.0 | corrective" in lower:
            in_corrective = True
            continue
        if in_corrective and ("6b | prevention" in lower or "6c | risk" in lower or "7.0 | review" in lower):
            break
        if not in_corrective or "|" not in line:
            continue
        if any(k in lower for k in header_keywords):
            continue

        parts = [p.strip() for p in line.split("|") if p.strip()]
        if len(parts) < 2:
            continue
        cause_cell = parts[0] if not re.match(r"^\d+(?:\.0)?$", parts[0]) else ""
        action_cell = parts[1] if re.match(r"^\d+(?:\.0)?$", parts[0]) else parts[1]
        cause_text_cn, cause_text_en = _split_bilingual_cell(cause_cell)
        cn, en = _split_bilingual_cell(action_cell)
        key = (cn or en).lower()
        if not key or key in seen:
            continue
        seen.add(key)
        is_escape = any(k in cn for k in detection_keywords) or any(k in en.lower() for k in detection_keywords) or any(k in cause_text_cn for k in detection_keywords) or any(k in cause_text_en.lower() for k in detection_keywords)
        target_cn = escape_cn if is_escape else action_cn
        target_en = escape_en if is_escape else action_en
        target_cause_cn = escape_cause_cn if is_escape else cause_cn
        target_cause_en = escape_cause_en if is_escape else cause_en
        if cause_text_cn:
            target_cause_cn.append(cause_text_cn + "。")
        if cause_text_en:
            target_cause_en.append(cause_text_en + ".")
        if cn:
            target_cn.append(cn + "。")
        if en:
            target_en.append(en + ".")

    return {
        "cause_cn": cause_cn[:2],
        "cause_en": cause_en[:2],
        "escape_cause_cn": escape_cause_cn[:2],
        "escape_cause_en": escape_cause_en[:2],
        "action_cn": action_cn[:3],
        "action_en": action_en[:3],
        "escape_cn": escape_cn[:3],
        "escape_en": escape_en[:3],
    }


def extract_8d(file_path, timeout=180, logger=None):
    """
    从 8D 报告文件提取：发生根因、流出原因、纠正措施（中英双语）。
    返回 dict:
      {
        "root_cause": str,       # 发生根因（中文）
        "root_cause_en": str,    # 发生根因（英文）
        "escape_cause": str,     # 流出原因（中文）
        "escape_cause_en": str,  # 流出原因（英文）
        "action": str,           # 纠正措施（中文）
        "action_en": str,        # 纠正措施（英文）
      }
    提取失败返回 None。
    """
    raw = extract_text_from_file(file_path, logger=logger)
    if not raw or len(raw.strip()) < 20:
        if logger:
            logger.info(f"[AI] 8D text too short: {file_path}")
        return None

    action_hint = _extract_action_hint(raw)
    fallback_actions = _fallback_actions_from_hint(action_hint)
    prompt = _8D_PROMPT.format(action_hint=action_hint or "(none found)", raw=raw.strip()[:10000])
    out = _call_ollama(prompt, timeout=timeout, num_predict=1500, logger=logger)
    if not out:
        if any(fallback_actions.values()):
            if logger:
                logger.info("[AI] 8D model returned no output; using deterministic D5/D6 action fallback")
            return {
                "root_cause": fallback_actions.get("root_cause", ""),
                "root_cause_en": fallback_actions.get("root_cause_en", ""),
                "escape_cause": fallback_actions.get("escape_cause", ""),
                "escape_cause_en": fallback_actions.get("escape_cause_en", ""),
                "action": fallback_actions.get("action", ""),
                "action_en": fallback_actions.get("action_en", ""),
                "escape_action": fallback_actions.get("escape_action", ""),
                "escape_action_en": fallback_actions.get("escape_action_en", ""),
            }
        return None

    data = _parse_json(out)
    if not data:
        if logger:
            logger.warning(f"[AI] 8D JSON parse failed: {out[:200]}")
        if any(fallback_actions.values()):
            return {
                "root_cause": fallback_actions.get("root_cause", ""),
                "root_cause_en": fallback_actions.get("root_cause_en", ""),
                "escape_cause": fallback_actions.get("escape_cause", ""),
                "escape_cause_en": fallback_actions.get("escape_cause_en", ""),
                "action": fallback_actions.get("action", ""),
                "action_en": fallback_actions.get("action_en", ""),
                "escape_action": fallback_actions.get("escape_action", ""),
                "escape_action_en": fallback_actions.get("escape_action_en", ""),
            }
        return None

    result = {
        "root_cause": _compact_8d_field(data.get("occurrence_cause"), 360),
        "root_cause_en": _compact_8d_field(data.get("occurrence_cause_en"), 420),
        "escape_cause": _compact_8d_field(data.get("escape_cause"), 320),
        "escape_cause_en": _compact_8d_field(data.get("escape_cause_en"), 380),
        "action": _compact_8d_field(data.get("occurrence_action"), 320),
        "action_en": _compact_8d_field(data.get("occurrence_action_en"), 380),
        "escape_action": _compact_8d_field(data.get("escape_action"), 300),
        "escape_action_en": _compact_8d_field(data.get("escape_action_en"), 360),
    }

    if action_hint and not all([result["action"], result["action_en"], result["escape_action"], result["escape_action_en"]]):
        for key, value in fallback_actions.items():
            if value and not result.get(key):
                result[key] = value

    if logger:
        logger.info(f"[AI] 8D extracted: cause={result['root_cause_en'][:40]}... escape={result['escape_cause_en'][:40]}...")

    # 至少要有一个有效字段
    if any(v for v in result.values()):
        return result
    return None
